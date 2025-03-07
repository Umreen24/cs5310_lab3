#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Tue Feb  9 21:42:29 2021

@author: umreenimam
"""

"""
IMPORTING PACKAGES
"""
import os
import math
import pandas as pd
from scipy.io import loadmat
import matplotlib.pyplot as plt
import seaborn as sns

from pyeeg import bin_power, spectral_entropy
from docx import Document
from sklearn.preprocessing import MinMaxScaler
from sklearn.model_selection import train_test_split
from sklearn.metrics import confusion_matrix, accuracy_score
from sklearn.naive_bayes import MultinomialNB

"""
Question 1 - Load the data in each of the Matlab files into a data frame.
"""
# Setting file names
os.chdir('/Users/umreenimam/Documents/Masters/Masters_Classes/CS_5310/Week_4/lab_ch4')
prefile = 'Pre(2).mat'
medfile = 'Med(2).mat'
postfile = 'Post(2).mat'

# Loading data 
predata = loadmat(prefile)
meddata = loadmat(medfile)
postdata = loadmat(postfile)

pre = predata['data']
med = meddata['data']
post = postdata['data']

# Transpose data and create dataframes 
pre_transposed = pd.DataFrame(pre.T)
med_transposed = pd.DataFrame(med.T)
post_transposed = pd.DataFrame(post.T)

"""
Question 2 - Compute the alpha PSIs and spectral entropies of 34 channels.
"""
# Creating function to calculate psis and spectral entropies
def get_psi_entropies(data):
    
    bands = [0.5, 4, 7, 12, 30, 100]
    fs = 125
    array_length = 1024
    
    cols = data.shape[1]
    rows = math.floor(data.shape[0] / array_length)
    alpha_psi_entropies = pd.DataFrame([])
    temp_data = pd.DataFrame()
    
    for x in range(cols):
        alpha_psis = []
        entropies = []
        temp_col_alpha = 'alpha_psi ' + str(x)
        temp_col_entropy = 'spectral_entropy ' + str(x)
        for y in range(rows):
            
            psis, power_ratios = bin_power(data.iloc[(y * array_length):((y + 1) * array_length), x], bands, fs)
            alpha_psis.append(psis[2])
            
            spec_entropies = spectral_entropy(data.iloc[(y * array_length):((y + 1) * array_length), x], bands, fs, power_ratios)
            entropies.append(spec_entropies)
            
        temp_data[temp_col_alpha] = alpha_psis
        temp_data[temp_col_entropy] = entropies
            
    alpha_psi_entropies = alpha_psi_entropies.append(temp_data, ignore_index = True)
    return alpha_psi_entropies

# Running function for each dataframe
pre_psi_entropy = get_psi_entropies(pre_transposed)
med_psi_entropy = get_psi_entropies(med_transposed)
post_psi_entropy = get_psi_entropies(post_transposed)

"""
Questions 3 & 4 - Create a list of brain activity labels corresponding 
to each data frame. Stack DataFrames vertically.
"""
pre_label = ["Pre"] * pre_psi_entropy.shape[0]
med_label = ["Med"] * med_psi_entropy.shape[0]
post_label = ["Post"] * post_psi_entropy.shape[0]

frames = [pre_psi_entropy, med_psi_entropy, post_psi_entropy]
combined_psi_entropies = pd.concat(frames, ignore_index = True)
combined_state_labels = pre_label + med_label + post_label

"""
Questions 5-7 - Create corrleation matrix of combined DataFrame, remove co-linearity,
print corrleation matrix of remaining channels.
"""
# Creating function to remove co-linearity 
def remove_co_linearity(data, neg_threshold):
    
    corr_mat = data.corr()
    row = corr_mat.shape[0]
    column = corr_mat.shape[1]
    
    correlated_features = []
    
    for x in range(row): 
        for y in range(column):
            if x == y:
                break
            if corr_mat.iloc[x, y] > abs(neg_threshold) or corr_mat.iloc[x, y] < neg_threshold:
                correlated_features.append(corr_mat.columns[x])
                break
    return corr_mat, correlated_features

# Running function to remove co-linearity 
corrleation_mat, cols_to_remove = remove_co_linearity(combined_psi_entropies, -0.9)
psi_entropies_remain = combined_psi_entropies.drop(columns = cols_to_remove, axis = 1)

# Creating correlation matrix of remaining values 
psi_entropies_reamin_corr = psi_entropies_remain.corr()

# Remaining values heatmap
sns.set_theme(style = "white")

plt.figure(figsize=(8,8))
color_map = sns.diverging_palette(230, 20, as_cmap = True)
sns.heatmap(corrleation_mat, annot = False, cmap = color_map, vmax = 1, 
            center = 0, square = True, linewidths = 0.1,
            cbar_kws = {"shrink": 0.75})
plt.title('Heat Map of Correlation Coefficient Matrix', fontsize = 18)
plt.xlabel('Column Number from the Data Frame', fontsize = 12)
plt.ylabel('Column Number from the Data Frame', fontsize = 12)
plt.savefig('fig1.png')
plt.show


plt.figure(figsize=(8,8))
color_map = sns.diverging_palette(230, 20, as_cmap = True)
sns.heatmap(psi_entropies_reamin_corr, annot = False, cmap = color_map, vmax = 1, 
            center = 0, square = True, linewidths = 0.5,
            cbar_kws = {"shrink": 0.75})
plt.title('Heat Map of Correlation Coefficient Matrix', fontsize = 18)
plt.xlabel('Column Number from the Data Frame', fontsize = 12)
plt.ylabel('Column Number from the Data Frame', fontsize = 12)
plt.savefig('fig2.png')
plt.show

"""
Question 8 - Normalize data using Min-Max method
"""
scaler = MinMaxScaler()
normalized_corr = scaler.fit_transform(psi_entropies_remain)
normalized_corr = pd.DataFrame(normalized_corr)

"""
Question 9 - Split data into training (80%) and testing (20%) sets.
Split labels accordingly.
"""
# Setting X & y variables for train, test
X = normalized_corr
y = combined_state_labels

# Splitting data into train and test models
X_train, X_test, y_train, y_test = train_test_split(
    X, 
    y, 
    random_state = 42,
    test_size = 0.2,
    stratify = y)

"""
Question 10 - Train a multinomial Naïve Bayes model 
using the training dataset. Set laplace = 1.
"""
# Training data using multinomial Naive Bayes model
nb = MultinomialNB(alpha = 1)
nb.fit(X_train, y_train)

"""
Question 11 - Use the Naïve Bayes model to predict the target 
feature of the testing dataset.
"""
# Predicting target feature of training dataset
predicted = nb.predict(X_test)

"""
Question 12 - Create a confusion matrix to compare the predicted state 
activities to the actual activities and compute the accuracy.
"""
# Creating confusion matrix 
confusion_mat = confusion_matrix(y_test, predicted)
print(confusion_mat)

# Computing accuracy
accuracy_rate = accuracy_score(y_test, predicted)
print(accuracy_rate)

"""
Generate Word document with results
"""
doc = Document()

doc.add_heading('Correlation Coefficient Matrix Before Removing Co-linearity', level = 1)

doc.add_picture('fig1.png')
doc.add_paragraph()
doc.add_page_break()

doc.add_heading('Correlation Coefficient Matrix After Removing Co-linearity', level = 1)

doc.add_picture('fig2.png')
doc.add_paragraph()

doc.add_heading('Confusion Matrix:', level = 1)

table = doc.add_table(rows = confusion_mat.shape[0] + 1, cols = confusion_mat.shape[1] + 1)
table.style = 'Medium Grid 3 Accent 3'

row = table.rows[0]
row.cells[1].text = 'Predicted Med'
row.cells[2].text = 'Predicted Post'
row.cells[3].text = 'Predicted Pre'

col = table.columns[0]
col.cells[1].text = 'Actual Med'
col.cells[2].text = 'Actual Post'
col.cells[3].text = 'Actual Pre'

for i in range(confusion_mat.shape[0]):
    for j in range(confusion_mat.shape[1]):
        table.cell(i + 1, j + 1).text = str(confusion_mat[i, j])

doc.add_paragraph()
doc.add_paragraph('Accuracy Rate: {}%'.format(round(accuracy_rate * 100, 1)))
doc.save('Chapter04-Lab-UI.docx')